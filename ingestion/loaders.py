
from typing import List
from langchain_core.documents import Document

import docx
import openpyxl
import zipfile


def _is_valid_docx(path: str) -> bool:
    try:
        with zipfile.ZipFile(path):
            return True
    except zipfile.BadZipFile:
        return False


def _load_license_from_plain_text(path: str) -> List[Document]:
   
    with open(path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    documents: List[Document] = []
    current_heading = ""
    current_text: List[str] = []

    def flush():
        if current_text:
            content = f"{current_heading}\n" + "\n".join(current_text)
            documents.append(
                Document(
                    page_content=content.strip(),
                    metadata={"source": "license_info", "section": current_heading or "General"},
                )
            )

    for raw_line in lines:
        text = raw_line.strip()
        if not text:
            continue

        if text.startswith("#"):
            flush()
            current_heading = text.lstrip("#").strip()
            current_text = []
        else:
            current_text.append(text)

    flush()
    return documents


def load_license_docx(path: str) -> List[Document]:
    

    if not _is_valid_docx(path):
        return _load_license_from_plain_text(path)

    doc = docx.Document(path)

    documents: List[Document] = []
    current_heading = ""
    current_text: List[str] = []

    def flush():
        if current_text:
            content = f"{current_heading}\n" + "\n".join(current_text)
            documents.append(
                Document(
                    page_content=content.strip(),
                    metadata={"source": "license_info", "section": current_heading or "General"},
                )
            )

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("Heading") or text.startswith("#"):
            flush()
            current_heading = text.lstrip("#").strip()
            current_text = []
        else:
            current_text.append(text)

    flush()
    return documents


def load_violations_xlsx(path: str) -> List[Document]:
    workbook = openpyxl.load_workbook(path, data_only=True)
    documents: List[Document] = []

    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(rows[0])]

        for row in rows[1:]:
            if row is None or all(v is None for v in row):
                continue
            row_dict = {headers[i]: (row[i] if i < len(row) else None) for i in range(len(headers))}
            content_lines = [f"{k}: {v}" for k, v in row_dict.items() if v is not None]
            content = "\n".join(content_lines)
            if content.strip():
                documents.append(
                    Document(
                        page_content=content,
                        metadata={"source": "violations", "sheet": sheet.title},
                    )
                )

    return documents


def load_all_documents(license_path: str, violations_path: str) -> List[Document]:
    docs = []
    docs.extend(load_license_docx(license_path))
    docs.extend(load_violations_xlsx(violations_path))
    return docs
